"""
doc_parser.py
Parses Word (.docx) and Excel (.xlsx) payroll documents.
"""

from app.core.input_handler import PaystubData
from app.model.gemma_client import GemmaClient


class DocParser:
    def __init__(self):
        self.gemma = GemmaClient()

    def parse_docx(self, file_path: str) -> PaystubData:
        print(f"Parsing Word doc: {file_path}")
        try:
            from docx import Document
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_text:
                        parts.append(" | ".join(row_text))
            structured = self.gemma.extract_paystub_fields("\n".join(parts))
            structured.confidence = 0.90
            return structured
        except Exception as e:
            print(f"Word parse error: {e}")
            return PaystubData(confidence=0.0, needs_review=True)

    def parse_xlsx(self, file_path: str) -> PaystubData:
        print(f"Parsing Excel: {file_path}")
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            parts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                parts.append(f"Sheet: {sheet}")
                for row in ws.iter_rows(values_only=True):
                    vals = [str(c) for c in row if c is not None]
                    if vals:
                        parts.append(" | ".join(vals))
            structured = self.gemma.extract_paystub_fields("\n".join(parts))
            structured.confidence = 0.90
            return structured
        except Exception as e:
            print(f"Excel parse error: {e}")
            return PaystubData(confidence=0.0, needs_review=True)